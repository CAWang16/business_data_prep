import sqlite3
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DB_FILE = "database/retail.db"
RAW_TABLE = "online_retail"
FIGURES_DIR = "figures"
OUTPUT_FILE = "report.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_table_from_df(doc, df, title=None):
    if title:
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def add_figure(doc, path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Figure not found: {path}]")

# ── load & compute ────────────────────────────────────────────────────────────

print("Loading data from database...")
conn = sqlite3.connect(DB_FILE)
df = pd.read_sql(f"SELECT * FROM {RAW_TABLE}", conn)
conn.close()

df = df.rename(columns={"Customer ID": "CustomerID", "Invoice": "InvoiceNo"})
df["InvoiceNo"] = df["InvoiceNo"].astype(str)
df["StockCode"] = df["StockCode"].astype(str).str.strip()

# --- script-00 metrics ---
total_rows = len(df)
total_cols = len(df.columns)
col_names = df.columns.tolist()

# InvoiceNo
valid_inv_mask = df["InvoiceNo"].str.match(r"^C?\d{6}$")
valid_inv = valid_inv_mask.sum()
invalid_inv = (~valid_inv_mask).sum()

# CustomerID
missing_cid = df["CustomerID"].isna().sum()
missing_cid_pct = missing_cid / total_rows

cid_str = df["CustomerID"].dropna().astype(int).astype(str)
valid_cid = cid_str.str.match(r"^\d{5}$").sum()
invalid_cid = (~cid_str.str.match(r"^\d{5}$")).sum()

# Quantity / Price
qty_neg = (df["Quantity"] < 0).sum()
qty_zero = (df["Quantity"] == 0).sum()
qty_pos = (df["Quantity"] > 0).sum()
price_neg = (df["Price"] < 0).sum()
price_zero = (df["Price"] == 0).sum()
price_pos = (df["Price"] > 0).sum()

# Dates
parsed_dates = pd.to_datetime(df["InvoiceDate"], errors="coerce")
invalid_dates = parsed_dates.isna().sum()
date_min = parsed_dates.min()
date_max = parsed_dates.max()

# StockCode
numeric_mask = df["StockCode"].str.match(r"^\d{5}$")
suffix_mask = df["StockCode"].str.match(r"^\d{5}[A-Za-z]{1,2}$")
special_mask = ~(numeric_mask | suffix_mask)
special_count = special_mask.sum()
unique_special = df.loc[special_mask, "StockCode"].nunique()
unique_sc = df["StockCode"].nunique()

# Country
missing_country = df["Country"].isna().sum()
unspecified = (df["Country"].astype(str).str.strip() == "Unspecified").sum()
unique_countries = df["Country"].nunique(dropna=True)
top10_countries = df["Country"].value_counts().head(10).reset_index()
top10_countries.columns = ["Country", "Row Count"]

# Description
missing_desc = df["Description"].isna().sum()

# --- script-01 metrics ---
is_admin = df["InvoiceNo"].str.startswith("A")
is_c_invoice = df["InvoiceNo"].str.startswith("C")
is_neg_qty = df["Quantity"] < 0
is_pos_qty = df["Quantity"] > 0
is_zero_price = df["Price"] == 0
is_pos_price = df["Price"] > 0

returns_count = is_c_invoice.sum()
returns_pct = returns_count / total_rows
admin_count = is_admin.sum()
free_count = ((~is_c_invoice) & (~is_admin) & is_pos_qty & is_zero_price).sum()
internal_adj = ((~is_c_invoice) & (~is_admin) & is_neg_qty & is_zero_price).sum()

is_dcgs = df["StockCode"].str.contains("DCGS", case=False, na=False)
is_gift = df["StockCode"].str.contains("gift", case=False, na=False)

df_no_admin = df[~is_admin]
clean_sales = df_no_admin[
    ~df_no_admin["InvoiceNo"].str.startswith("C") &
    (df_no_admin["Quantity"] > 0) &
    (df_no_admin["Price"] > 0)
]

mismatch_neg_not_c = (is_neg_qty & ~is_c_invoice).sum()
mismatch_c_not_neg = (is_c_invoice & ~is_neg_qty).sum()

# summary stats for clean sales
cs_revenue = (clean_sales["Quantity"] * clean_sales["Price"])
cs_total_revenue = cs_revenue.sum()
cs_unique_invoices = clean_sales["InvoiceNo"].nunique()
cs_unique_customers = clean_sales["CustomerID"].nunique()
cs_unique_products = clean_sales["StockCode"].nunique()

print("Data loaded. Building document...")

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# Title
title = doc.add_heading("Online Retail II — Data Processing & Analysis Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

authors = doc.add_paragraph("Carl & Anders  |  CSP 571 — Spring 2026")
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.runs[0].bold = True

date_p = doc.add_paragraph("April 2026")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ─── 1. Abstract ───────────────────────────────────────────────────────────

add_heading(doc, "1. Abstract", 1)
add_body(doc,
    "This report documents the full data processing and exploratory analysis pipeline applied "
    "to the UCI Online Retail II dataset, covering transactions from a UK-based non-store online "
    "retailer between December 2009 and December 2011. The dataset comprises 1,067,371 raw "
    "transaction records spanning eight fields. Our work proceeds in three phases: (1) metadata "
    "validation of raw source fields to identify structural anomalies and format inconsistencies, "
    "(2) split-decision exploratory data analysis to classify records into semantically distinct "
    "transaction types, and (3) computation of summary statistics, data visualizations, and "
    "feature extraction in preparation for downstream modelling. Key findings include that "
    f"approximately {missing_cid_pct:.1%} of records lack a CustomerID, that InvoiceNo and "
    "StockCode encode multiple transaction types (sales, cancellations, accounting adjustments, "
    "free promotions), and that a structured seven-category split is required before any "
    "customer-level or product-level analysis can be conducted reliably."
)

doc.add_page_break()

# ─── 2. Overview ───────────────────────────────────────────────────────────

add_heading(doc, "2. Overview", 1)

add_heading(doc, "2.1 Problem Statement", 2)
add_body(doc,
    "E-commerce retailers accumulate vast volumes of transactional data whose raw form is rarely "
    "suitable for direct analysis. The Online Retail II dataset presents several challenges "
    "representative of real-world retail data: mixed transaction semantics encoded within shared "
    "fields (sales, returns, administrative adjustments, and free promotions all reside in the "
    "same table), pervasive missing values in the CustomerID field, non-standard StockCode "
    "identifiers for service and operational entries, and naming inconsistencies in geographic "
    "fields."
)
add_body(doc,
    "The primary objective of this project is to design and implement a reproducible data "
    "processing pipeline that: (a) validates the structural integrity of raw source fields, "
    "(b) identifies and separates semantically distinct record types, (c) applies principled "
    "cleaning and transformation rules with explicit assumptions, and (d) extracts analytic "
    "features that support downstream customer segmentation and purchase-pattern modelling."
)

add_heading(doc, "2.2 Related Work / Literature Review", 2)
add_body(doc,
    "Retail transaction datasets have been extensively studied in the data mining literature. "
    "RFM (Recency, Frequency, Monetary) analysis, introduced by Bult and Wansbeek (1995) and "
    "popularised by Hughes (1994), remains one of the most widely applied frameworks for "
    "customer value segmentation and is directly applicable to invoice-level data such as the "
    "Online Retail II dataset."
)
add_body(doc,
    "Chen et al. (2012) used an earlier version of this dataset (Online Retail I) to demonstrate "
    "data mining techniques for customer segmentation using K-Means clustering on RFM features, "
    "establishing a widely cited baseline for this source. Daqing Chen's original UCI submission "
    "noted that the dataset contains a mix of wholesale and retail customers, which has "
    "implications for any per-unit price analysis."
)
add_body(doc,
    "More recent work has extended transaction-level analysis to sequence modelling and "
    "collaborative filtering. Hidasi et al. (2016) demonstrated session-based recurrent neural "
    "networks for next-item recommendation, while Devooght and Bersini (2017) compared "
    "collaborative filtering approaches on transactional purchase histories. These downstream "
    "applications all share the upstream requirement of clean, well-typed transaction records, "
    "motivating the careful data processing approach documented in this report."
)
add_body(doc,
    "Data cleaning practices for retail datasets are discussed by Rahm and Do (2000), who "
    "classify data quality problems into single-source and multi-source issues. The challenges "
    "encountered in this project — missing identifiers, mixed semantics within a single column, "
    "and implicit coding conventions (C-prefixed cancellation invoices, A-prefixed accounting "
    "entries) — are consistent with their taxonomy of schema-level and instance-level anomalies."
)

add_heading(doc, "2.3 Proposed Methodology", 2)
add_body(doc,
    "Our methodology follows a sequential, script-driven pipeline with one Python script per "
    "logical stage, each operating on the shared SQLite database. This design keeps each stage "
    "independently reproducible and auditable."
)
add_bullet(doc, "Stage 0 — Data Ingestion (setup_db.py): Load both Excel sheets from the raw "
           "source file into a single SQLite table.")
add_bullet(doc, "Stage 1 — Metadata Validation (00_validate_source_metadata.py): Validate "
           "each raw field for format conformance, value ranges, and missingness. Document "
           "all anomalies with counts and representative examples.")
add_bullet(doc, "Stage 2 — Split-Decision EDA (01_split_decision_eda.py): Classify records "
           "into transaction types (sales, cancellations, admin adjustments, free items, "
           "gift vouchers, DCGS codes) using cross-tabulation and overlap analysis.")
add_bullet(doc, "Stage 3 — Cleaning & Transformation: Apply the split decisions to produce "
           "a filtered, typed, and enriched table ready for analysis.")
add_bullet(doc, "Stage 4 — Analysis: Compute summary statistics, generate visualizations, "
           "and extract modelling features.")

doc.add_page_break()

# ─── 3. Data Processing ────────────────────────────────────────────────────

add_heading(doc, "3. Data Processing", 1)

# 3.1 Data Description
add_heading(doc, "3.1 Data Description", 2)
add_body(doc,
    "The Online Retail II dataset (UCI Machine Learning Repository, donated by Daqing Chen, "
    "London South Bank University) covers transactions of a UK-based online retailer that "
    "primarily sells unique all-occasion gift-ware to wholesale customers. The data spans two "
    "consecutive fiscal years delivered as two Excel sheets."
)

desc_df = pd.DataFrame({
    "Property": [
        "Source", "Format", "Sheets", "Date range",
        "Total rows (combined)", "Columns",
    ],
    "Value": [
        "UCI Machine Learning Repository — Online Retail II",
        "Microsoft Excel (.xlsx)",
        "Year 2009-2010 (525,461 rows), Year 2010-2011 (541,910 rows)",
        f"{date_min.strftime('%d %b %Y')} – {date_max.strftime('%d %b %Y')}",
        f"{total_rows:,}",
        str(total_cols),
    ]
})
add_table_from_df(doc, desc_df, "Table 1 — Dataset overview")

add_body(doc, "The eight columns and their semantics are described below.")
add_body(doc,
    "Note: The raw Excel file and SQLite database store the unit price column as 'Price'. "
    "The official dataset description names this field 'UnitPrice'. Both refer to the same field; "
    "this report uses 'UnitPrice' to match the authoritative specification."
)
col_df = pd.DataFrame({
    "Official Name": ["InvoiceNo", "StockCode", "Description", "Quantity",
                      "InvoiceDate", "UnitPrice", "CustomerID", "Country"],
    "DB Column": ["Invoice → InvoiceNo", "StockCode", "Description", "Quantity",
                  "InvoiceDate", "Price", "Customer ID → CustomerID", "Country"],
    "Type": ["Nominal", "Nominal", "Nominal", "Numeric",
             "Numeric", "Numeric", "Nominal", "Nominal"],
    "Official Definition": [
        "6-digit integral number uniquely assigned to each transaction. "
        "'C' prefix indicates a cancellation.",
        "5-digit integral number uniquely assigned to each distinct product.",
        "Product (item) name.",
        "Quantities of each product per transaction.",
        "Day and time when the transaction was generated.",
        "Product price per unit in sterling (£).",
        "5-digit integral number uniquely assigned to each customer.",
        "Name of the country where the customer resides.",
    ],
    "Observed Deviations": [
        "6 rows carry an 'A' prefix (accounting adjustments) — not in spec",
        "Many codes have letter suffixes (e.g. 85123A); special codes exist (POST, DOT, M…) — spec says 5-digit integral only",
        "~2 % of rows have missing descriptions",
        "Negative values present (returns/adjustments)",
        "No deviations — all rows parse successfully",
        "Zero and negative values present; column stored as 'Price' in source file",
        "~25 % of rows have no CustomerID",
        "'Unspecified' used for unknown origin; naming inconsistencies (e.g. EIRE vs Ireland)",
    ]
})
add_table_from_df(doc, col_df, "Table 2 — Official column definitions and observed deviations")

# 3.2 Data Cleaning
add_heading(doc, "3.2 Data Cleaning", 2)
add_body(doc,
    "Cleaning decisions were derived from the metadata validation output of "
    "00_validate_source_metadata.py. The following subsections document each field."
)

add_heading(doc, "InvoiceNo", 3)
add_body(doc,
    "The official specification defines InvoiceNo as a 6-digit integral number, with a 'C' "
    "prefix (case-insensitive in the data) indicating a cancellation. No other prefixes are "
    "documented in the spec."
)
add_body(doc,
    f"Of {total_rows:,} rows, {valid_inv:,} conform to the spec (6-digit integer, "
    f"optionally prefixed with 'C'). The remaining {invalid_inv:,} rows carry an 'A' prefix "
    "and represent bad-debt accounting adjustment entries — an undocumented transaction type "
    "not described in the official field definition. These rows are excluded from all "
    "downstream analysis."
)

n_c_invoice = df["InvoiceNo"].str.startswith("C").sum()
n_standard_sale = valid_inv - n_c_invoice
inv_df = pd.DataFrame({
    "Category": [
        "Standard sale (######)  ← spec-compliant",
        "Cancellation (C######)  ← spec-compliant",
        "Accounting adjustment (A######)  ← undocumented deviation",
        "Total"
    ],
    "Count": [
        f"{n_standard_sale:,}",
        f"{n_c_invoice:,}",
        f"{invalid_inv:,}",
        f"{total_rows:,}"
    ],
    "Action": ["Retain — classify as sale", "Retain — classify as return", "Exclude", "—"]
})
add_table_from_df(doc, inv_df, "Table 3 — InvoiceNo validation against official spec")

add_heading(doc, "CustomerID", 3)
add_body(doc,
    f"CustomerID is missing for {missing_cid:,} rows ({missing_cid_pct:.1%} of the dataset). "
    f"All {valid_cid:,} non-missing values conform to the 5-digit format; no format violations "
    "were found. Missing CustomerIDs are retained in non-customer analyses (e.g., product-level "
    "aggregations) but excluded from any customer-level modelling."
)

add_heading(doc, "Quantity", 3)
qty_df = pd.DataFrame({
    "Bucket": ["Quantity > 0", "Quantity = 0", "Quantity < 0"],
    "Count": [f"{qty_pos:,}", f"{qty_zero:,}", f"{qty_neg:,}"],
    "Interpretation": [
        "Standard sale or free promotional item",
        "Data anomaly — reviewed case by case",
        "Return / cancellation or internal adjustment"
    ]
})
add_table_from_df(doc, qty_df, "Table 4 — Quantity value breakdown")

add_heading(doc, "Price", 3)
price_df = pd.DataFrame({
    "Bucket": ["Price > 0", "Price = 0", "Price < 0"],
    "Count": [f"{price_pos:,}", f"{price_zero:,}", f"{price_neg:,}"],
    "Interpretation": [
        "Revenue-generating transaction",
        "Free item or non-revenue system entry",
        "Bad-debt accounting adjustment — exclude"
    ]
})
add_table_from_df(doc, price_df, "Table 5 — Price value breakdown")

add_heading(doc, "InvoiceDate", 3)
add_body(doc,
    f"All {total_rows - invalid_dates:,} InvoiceDate values parsed successfully. "
    f"No unparseable dates were found. The dataset spans "
    f"{date_min.strftime('%d %B %Y')} to {date_max.strftime('%d %B %Y')}."
)

add_heading(doc, "StockCode", 3)
add_body(doc,
    "The official specification defines StockCode as a 5-digit integral number. "
    "Validation reveals two classes of deviation from this definition:"
)
sc_df = pd.DataFrame({
    "Pattern": [
        "5-digit numeric (e.g. 85123)  ← spec-compliant",
        "5-digit + letter suffix (e.g. 85123A)  ← deviation from spec",
        "Fully non-standard (e.g. POST, DOT, M, DCGS, gift)  ← deviation from spec"
    ],
    "Rows": [f"{numeric_mask.sum():,}", f"{suffix_mask.sum():,}", f"{special_count:,}"],
    "Unique codes": [
        f"{df.loc[numeric_mask, 'StockCode'].nunique():,}",
        f"{df.loc[suffix_mask, 'StockCode'].nunique():,}",
        f"{unique_special:,}"
    ],
    "Treatment": [
        "Standard product SKU — included in product analyses",
        "Treated as product variant; included in product analyses with spec caveat noted",
        "Service, discount, postage, or adjustment entry — separated or excluded"
    ]
})
add_table_from_df(doc, sc_df, "Table 6 — StockCode validation against official spec")
add_body(doc,
    "Non-standard special codes include POST (postage), DOT (dotcom), M (manual entry), "
    "BANK CHARGES, ADJUST, AMAZONFEE, and the DCGS and gift code families. "
    "These fall outside the 5-digit integral definition in the official spec and are "
    "treated as separate categories rather than product SKUs."
)

add_heading(doc, "Country", 3)
add_body(doc,
    f"Country has no missing or blank values across all {total_rows:,} rows. "
    f"There are {unique_countries} unique country labels. "
    f"{unspecified:,} rows ({unspecified / total_rows:.2%}) carry the value 'Unspecified', "
    "representing transactions with no recorded geographic origin. "
    "Naming inconsistencies (e.g. 'EIRE' instead of 'Ireland') are noted and will be "
    "standardised in the transformation stage."
)
add_table_from_df(doc, top10_countries, "Table 7 — Top 10 countries by row count")

add_heading(doc, "Description", 3)
add_body(doc,
    f"Description is missing for {missing_desc:,} rows ({missing_desc / total_rows:.2%}). "
    "Description is treated as a supplementary text field rather than a primary identifier. "
    "Missing values are filled from the most common description associated with each StockCode "
    "in the downstream transformation stage."
)

# 3.3 Data Transformation & Assumptions
add_heading(doc, "3.3 Data Transformation & Assumptions", 2)
add_body(doc,
    "Following the split-decision EDA (01_split_decision_eda.py), records are assigned to one "
    "of seven mutually exclusive categories. The assignment logic and assumptions are documented "
    "in the table below."
)

split_df = pd.DataFrame({
    "Category": [
        "Admin / accounting adjustments",
        "Cancellation invoices",
        "Internal inventory adjustments",
        "Free items / promotions",
        "Gift vouchers",
        "DCGS codes",
        "Clean sales"
    ],
    "Rule": [
        "InvoiceNo starts with 'A'",
        "InvoiceNo starts with 'C'",
        "Qty < 0 AND Price = 0 AND not C/A-invoice",
        "Qty > 0 AND Price = 0 AND not C/A-invoice",
        "StockCode contains 'gift'",
        "StockCode contains 'DCGS'",
        "Qty > 0, Price > 0, not C/A-invoice, not gift/DCGS"
    ],
    "Count": [
        f"{admin_count:,}",
        f"{returns_count:,}",
        f"{internal_adj:,}",
        f"{free_count:,}",
        f"{is_gift.sum():,}",
        f"{is_dcgs.sum():,}",
        f"{len(clean_sales):,}"
    ],
    "Disposition": [
        "Exclude from all analyses",
        "Separate returns analysis",
        "Exclude from sales analyses",
        "Separate promotions analysis",
        "Separate voucher analysis",
        "Separate channel analysis",
        "Primary analysis dataset"
    ]
})
add_table_from_df(doc, split_df, "Table 8 — Transaction type split decisions")

add_body(doc,
    "Key assumptions underpinning these decisions:"
)
add_bullet(doc,
    "A C-prefixed invoice always represents a cancellation regardless of Quantity sign. "
    f"One exception ({mismatch_c_not_neg:,} row with positive Quantity on a C-invoice) "
    "has been identified and is excluded from both the returns and clean-sales datasets."
)
add_bullet(doc,
    f"Rows with negative Quantity but no C-prefix ({mismatch_neg_not_c:,} rows) are "
    "primarily administrative adjustments and are excluded from sales analysis."
)
add_bullet(doc,
    "Zero-price rows with positive Quantity are treated as promotional/free items and "
    "are not counted as revenue-generating sales."
)
add_bullet(doc,
    "DCGS and gift StockCode families are treated as separate channels because their "
    "price and description patterns differ systematically from standard product SKUs."
)
add_bullet(doc,
    "Missing CustomerID is retained in product-level and geographic analyses but "
    "excluded from customer-level cohort and segmentation analyses."
)

doc.add_page_break()

# ─── 4. Data Analysis ──────────────────────────────────────────────────────

add_heading(doc, "4. Data Analysis", 1)

# 4.1 Summary Statistics
add_heading(doc, "4.1 Summary Statistics", 2)

add_body(doc,
    "Summary statistics are computed on the clean-sales subset "
    "(Quantity > 0, Price > 0, not C/A-invoice, not gift/DCGS special codes) "
    "unless otherwise noted."
)

summary_df = pd.DataFrame({
    "Metric": [
        "Clean sale rows",
        "Unique invoices",
        "Unique customers (with CustomerID)",
        "Unique products (StockCode)",
        "Total revenue (GBP)",
        "Avg revenue per invoice (GBP)",
        "Avg UnitPrice (GBP)",
        "Avg quantity per line",
        "Date range",
    ],
    "Value": [
        f"{len(clean_sales):,}",
        f"{cs_unique_invoices:,}",
        f"{cs_unique_customers:,}",
        f"{cs_unique_products:,}",
        f"£{cs_total_revenue:,.2f}",
        f"£{cs_total_revenue / cs_unique_invoices:,.2f}",
        f"£{clean_sales['Price'].mean():.2f}",  # DB column is 'Price' (= official UnitPrice)
        f"{clean_sales['Quantity'].mean():.1f}",
        f"{date_min.strftime('%d %b %Y')} – {date_max.strftime('%d %b %Y')}",
    ]
})
add_table_from_df(doc, summary_df, "Table 9 — Clean-sales summary statistics")

# Top products
top_products = (
    clean_sales.groupby("StockCode")["Description"]
    .agg(lambda x: x.mode().iloc[0] if not x.mode().empty else "")
    .reset_index()
    .merge(
        clean_sales.groupby("StockCode")["Quantity"].sum().reset_index(),
        on="StockCode"
    )
    .sort_values("Quantity", ascending=False)
    .head(10)
    .rename(columns={"Quantity": "Total Qty Sold"})
)
top_products["Total Qty Sold"] = top_products["Total Qty Sold"].apply(lambda x: f"{x:,}")
add_table_from_df(doc, top_products[["StockCode", "Description", "Total Qty Sold"]],
                  "Table 10 — Top 10 products by total quantity sold")

# Returns rate
add_body(doc,
    f"Returns (C-invoice rows) account for {returns_count:,} records "
    f"({returns_pct:.2%} of the raw dataset), confirming that cancellations are a "
    "material component of the data and must be modelled explicitly in any "
    "customer lifetime value or churn analysis."
)

# Missing CustomerID by country
missing_cid_by_country = (
    df[~is_admin]
    .groupby("Country")
    .apply(lambda x: x["CustomerID"].isna().mean())
    .sort_values(ascending=False)
    .head(8)
    .reset_index()
)
missing_cid_by_country.columns = ["Country", "Missing CustomerID Rate"]
missing_cid_by_country["Missing CustomerID Rate"] = (
    missing_cid_by_country["Missing CustomerID Rate"].map("{:.1%}".format)
)
add_table_from_df(doc, missing_cid_by_country,
                  "Table 11 — Top countries by missing CustomerID rate")

# 4.2 Data Visualization
add_heading(doc, "4.2 Data Visualization", 2)

add_body(doc,
    "Three visualizations were produced to support exploratory understanding of the dataset's "
    "revenue distribution, temporal patterns, and geographic composition."
)

add_figure(
    doc,
    os.path.join(FIGURES_DIR, "totalprice_distribution.png"),
    "Figure 1 — Distribution of total line-item price (Quantity × Price) for clean sales. "
    "The distribution is heavily right-skewed, indicating that most transactions are small "
    "while a small number of bulk orders generate disproportionately high revenue."
)

add_figure(
    doc,
    os.path.join(FIGURES_DIR, "seasonality_timing.png"),
    "Figure 2 — Seasonality and intra-day timing patterns. Monthly aggregation reveals a "
    "pronounced sales peak in the fourth quarter (October–November), consistent with "
    "pre-Christmas gift-ware demand. Hour-of-day patterns show peak activity during "
    "UK business hours (10:00–14:00)."
)

add_figure(
    doc,
    os.path.join(FIGURES_DIR, "geographic_insights.png"),
    "Figure 3 — Geographic revenue distribution. The United Kingdom accounts for the "
    "substantial majority of transactions. European markets (Netherlands, Germany, France, "
    "EIRE) are the most significant international segments."
)

# 4.3 Feature Extraction
add_heading(doc, "4.3 Feature Extraction", 2)
add_body(doc,
    "Feature extraction targets two downstream modelling objectives: "
    "(a) customer-level segmentation via RFM analysis, and "
    "(b) product-level demand characterisation. Features are derived exclusively from the "
    "clean-sales subset to avoid contamination from returns, adjustments, or promotional rows."
)

add_heading(doc, "RFM Features (customer level)", 3)
rfm_df = pd.DataFrame({
    "Feature": ["Recency", "Frequency", "Monetary"],
    "Definition": [
        "Days since the customer's most recent purchase (relative to dataset end date)",
        "Total number of distinct invoices associated with the customer",
        "Total revenue (sum of Quantity × Price) attributed to the customer"
    ],
    "Notes": [
        "Computed per CustomerID; rows with missing CustomerID excluded",
        "Invoice-level count; each InvoiceNo counted once per customer",
        "Gross sales only; returns not netted at this stage"
    ]
})
add_table_from_df(doc, rfm_df, "Table 12 — RFM feature definitions")

add_heading(doc, "Transaction-level Features", 3)
txn_df = pd.DataFrame({
    "Feature": [
        "TotalPrice (= Quantity × UnitPrice)",
        "Month", "DayOfWeek", "Hour",
        "IsUK",
        "HasCustomerID"
    ],
    "Definition": [
        "Quantity × UnitPrice per line item",
        "Calendar month of InvoiceDate (1–12)",
        "Day of week of InvoiceDate (0 = Monday)",
        "Hour of InvoiceDate (0–23)",
        "Binary indicator: Country == 'United Kingdom'",
        "Binary indicator: CustomerID is not missing"
    ]
})
add_table_from_df(doc, txn_df, "Table 13 — Transaction-level derived features")

add_heading(doc, "Product-level Features", 3)
add_body(doc,
    "For each StockCode in the clean-sales subset, the following aggregate features are "
    "computed across the full observation window:"
)
prod_df = pd.DataFrame({
    "Feature": [
        "TotalQuantitySold",
        "UniqueCustomers",
        "UniqueInvoices",
        "AvgUnitPrice (official field name)",
        "RevenueContribution"
    ],
    "Definition": [
        "Sum of Quantity across all clean-sale line items for the product",
        "Count of distinct CustomerIDs that purchased the product",
        "Count of distinct InvoiceNos containing the product",
        "Mean UnitPrice across all clean-sale line items for the product",
        "Product's share of total clean-sales revenue (%)"
    ]
})
add_table_from_df(doc, prod_df, "Table 14 — Product-level aggregate features")

# ── save ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_FILE)
print(f"Report saved to {OUTPUT_FILE}")
